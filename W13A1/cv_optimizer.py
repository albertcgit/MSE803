#!/usr/bin/env python3
import argparse
import json
import os
import re
import subprocess
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pypdf import PdfReader
from google import genai
from google.genai import types

# Paste your Gemini API key here (keep this file private / gitignored)
API_KEY = "REPLACE-WITH-YOUR-GEMINI-KEY"


# Contact details extracted from the CV header
@dataclass
class ContactInfo:
    name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    other_links: List[str] = field(default_factory=list)


# A labeled section of the CV (Experience, Education, etc.)
@dataclass
class CVSection:
    title: str
    content: List[str] = field(default_factory=list)


# Full parsed representation of a CV
@dataclass
class ParsedCV:
    raw_text: str
    contact: ContactInfo
    summary: List[str] = field(default_factory=list)
    sections: List[CVSection] = field(default_factory=list)


# Feedback result from either the LLM or the rule-based engine
@dataclass
class Feedback:
    overall_score: int
    strengths: List[str]
    weaknesses: List[str]
    recommendations: List[str]
    rewritten_summary: Optional[str] = None
    rewritten_bullets: Optional[dict] = None
    raw_llm_response: Optional[str] = None


SECTION_HEADERS = [
    "summary", "objective", "profile", "about",
    "experience", "work experience", "professional experience", "employment history",
    "education", "academic background",
    "skills", "technical skills", "core competencies",
    "projects", "certifications", "certificates",
    "awards", "achievements", "publications", "languages",
    "volunteer", "volunteering", "interests", "references",
]

EMAIL_RE = re.compile(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+")
PHONE_RE = re.compile(r"(\+?\d[\d\-\s().]{7,}\d)")
LINKEDIN_RE = re.compile(r"(linkedin\.com/\S+)", re.IGNORECASE)
URL_RE = re.compile(r"(https?://\S+|(?:www\.)\S+\.\S+)", re.IGNORECASE)

ACTION_VERBS = {
    "achieved", "improved", "led", "managed", "built", "created", "designed",
    "developed", "implemented", "launched", "reduced", "increased", "optimized",
    "streamlined", "delivered", "spearheaded", "negotiated", "automated",
    "analyzed", "coordinated", "directed", "established", "generated",
    "initiated", "mentored", "resolved", "trained", "transformed",
}

WEAK_PHRASES = [
    "responsible for", "duties included", "worked on", "helped with",
    "in charge of", "tasked with", "various", "team player", "hard worker",
    "detail oriented", "detail-oriented", "go-getter", "think outside the box",
]

BUZZWORDS = ["synergy", "leverage", "dynamic", "results-driven", "self-starter", "proactive"]

LLM_SYSTEM_PROMPT = """You are an expert career coach and professional CV/resume reviewer.
You give constructive, specific, and actionable feedback on CVs, covering:
1. CONTENT - impact of bullet points, quantified achievements, relevance, keyword optimization for ATS.
2. STRUCTURE - section organization, ordering, completeness (summary, experience, education, skills).
3. PRESENTATION - clarity, conciseness, consistency, tone, professionalism.

You will be given only the individual ACHIEVEMENT/RESPONSIBILITY BULLET LINES extracted from the
CV (job titles, company names, dates, and section headers have already been stripped out - do not
expect them and do not try to reconstruct them).

Always reply with STRICT JSON only, matching this schema (no markdown fences, no extra text):
{
  "overall_score": <integer 0-100>,
  "strengths": [<string>, ...],
  "weaknesses": [<string>, ...],
  "recommendations": [<string>, ...],
  "rewritten_summary": <string, a punchy improved 2-3 sentence professional summary>,
  "rewritten_bullets": {
      "<exact original bullet line as given to you>": "<the same bullet, rewritten>",
      ...
  }
}

Rules for rewritten_bullets:
- Only include a key for a bullet if you are actually improving it (weak verb, no quantified
  result, passive phrasing, wordiness). Skip bullets that are already strong.
- The key MUST be copied character-for-character from the input bullet list so it can be matched
  back up - do not paraphrase the key, only the value.
- Never invent employers, dates, numbers, or achievements not implied by the original text.
- Keep each rewritten bullet to a single line, starting with a strong action verb.
"""

ACCENT_COLOR = RGBColor(0x1F, 0x4E, 0x79)


# Reads a CV from .txt, .docx, or .pdf into plain text
def read_cv_file(path: str) -> str:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"CV file not found: {path}")
    suffix = p.suffix.lower()

    if suffix in (".txt", ".md"):
        return p.read_text(encoding="utf-8", errors="ignore")

    if suffix == ".docx":
        d = docx.Document(str(p))
        parts = [para.text for para in d.paragraphs if para.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if suffix == ".pdf":
        reader = PdfReader(str(p))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    raise ValueError(f"Unsupported file type: {suffix}. Use .txt, .docx, or .pdf")


# Detects whether a line is a section header
def _looks_like_header(line: str) -> Optional[str]:
    clean = line.strip().strip(":").strip()
    low = clean.lower()
    for header in SECTION_HEADERS:
        if low == header or low.replace(" ", "") == header.replace(" ", ""):
            return clean
    if 0 < len(clean) <= 35 and clean == clean.upper() and any(c.isalpha() for c in clean):
        return clean
    return None


# Parses raw CV text into contact info, summary, and sections
def parse_cv(raw_text: str) -> ParsedCV:
    lines = [l.rstrip() for l in raw_text.splitlines()]
    lines = [l for l in lines if l.strip() != ""]

    contact = ContactInfo()
    body_start_idx = 0

    head_block = "\n".join(lines[:10])
    email_match = EMAIL_RE.search(head_block)
    phone_match = PHONE_RE.search(head_block)
    linkedin_match = LINKEDIN_RE.search(head_block)
    if email_match:
        contact.email = email_match.group(0)
    if phone_match:
        contact.phone = phone_match.group(0).strip()
    if linkedin_match:
        contact.linkedin = linkedin_match.group(0)

    for url in URL_RE.findall(head_block):
        if "linkedin" not in url.lower():
            contact.other_links.append(url)

    if lines:
        first = lines[0].strip()
        if not EMAIL_RE.search(first) and not _looks_like_header(first):
            contact.name = first
            body_start_idx = 1

    for l in lines[1:5]:
        if "," in l and "@" not in l and len(l) < 60 and not URL_RE.search(l):
            contact.location = l.strip()
            break

    sections: List[CVSection] = []
    summary_lines: List[str] = []
    current: Optional[CVSection] = None

    for line in lines[body_start_idx:]:
        header = _looks_like_header(line)
        if header:
            current = CVSection(title=header)
            sections.append(current)
            continue
        if current is None:
            summary_lines.append(line)
        else:
            current.content.append(line)

    def _is_contact_only_line(line: str) -> bool:
        stripped = line.strip()
        if EMAIL_RE.search(stripped) or PHONE_RE.search(stripped) or LINKEDIN_RE.search(stripped):
            return True
        return False

    summary_lines = [l for l in summary_lines if not _is_contact_only_line(l)]

    return ParsedCV(
        raw_text=raw_text,
        contact=contact,
        summary=_reflow_lines(summary_lines),
        sections=[CVSection(title=s.title, content=_reflow_lines(s.content)) for s in sections],
    )


JOB_HEADER_RE = re.compile(r"(\||·).*\(.*\d{4}.*\)")


# Merges PDF-wrapped continuation fragments back into full logical lines
def _reflow_lines(lines: List[str]) -> List[str]:
    merged: List[str] = []
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        is_bullet = line.startswith(("-", "•", "*"))
        is_job_header = bool(JOB_HEADER_RE.search(line))
        if not merged:
            merged.append(line)
            continue
        prev = merged[-1]
        prev_ends_sentence = prev.rstrip().endswith((".", ":", "!", "?", ")"))
        if is_bullet or is_job_header or prev_ends_sentence:
            merged.append(line)
        else:
            merged[-1] = prev + " " + line
    return merged


# Checks whether a line is a quantifiable number
def _has_quantified_metric(line: str) -> bool:
    return bool(re.search(r"\d", line))


# Offline heuristic feedback engine, used when no LLM call is made
def rule_based_feedback(cv: ParsedCV) -> Feedback:
    strengths: List[str] = []
    weaknesses: List[str] = []
    recommendations: List[str] = []
    score = 100

    if not cv.contact.email:
        weaknesses.append("No email address detected - recruiters need a direct way to reach you.")
        recommendations.append("Add a professional email address near the top of the CV.")
        score -= 10
    else:
        strengths.append("Email address is present and easy to find.")

    if not cv.contact.phone:
        weaknesses.append("No phone number detected.")
        recommendations.append("Include a phone number in the header/contact block.")
        score -= 5

    if not cv.contact.linkedin:
        recommendations.append("Consider adding a LinkedIn profile URL to strengthen your professional presence.")
        score -= 3
    else:
        strengths.append("LinkedIn profile is included.")

    summary_text = " ".join(cv.summary).strip()
    if not summary_text:
        weaknesses.append("Missing a professional summary/profile statement at the top.")
        recommendations.append("Add a 2-3 sentence summary highlighting your role, years of experience, and top strengths.")
        score -= 10
    elif len(summary_text.split()) < 15:
        weaknesses.append("Professional summary is very short and doesn't convey much value.")
        recommendations.append("Expand the summary to 2-3 sentences covering your specialty, key skills, and a standout achievement.")
        score -= 5
    else:
        strengths.append("Includes a professional summary that frames the rest of the CV.")

    section_titles_lower = [s.title.lower() for s in cv.sections]
    expected = {
        "experience": ["experience", "employment history", "work experience", "professional experience"],
        "education": ["education", "academic background"],
        "skills": ["skills", "technical skills", "core competencies"],
    }
    for key, variants in expected.items():
        if not any(v in " ".join(section_titles_lower) for v in variants):
            weaknesses.append(f"No clearly labeled '{key.title()}' section was found.")
            recommendations.append(f"Add a clearly labeled '{key.title()}' section so recruiters and ATS software can find it.")
            score -= 8
        else:
            strengths.append(f"'{key.title()}' section is present and clearly labeled.")

    all_bullets = [line for s in cv.sections for line in s.content]
    weak_hits = 0
    quantified_hits = 0
    action_verb_hits = 0

    for line in all_bullets:
        low = line.lower()
        if any(p in low for p in WEAK_PHRASES):
            weak_hits += 1
        if _has_quantified_metric(line):
            quantified_hits += 1
        first_word = re.sub(r"[^a-zA-Z]", "", low.split()[0]) if low.split() else ""
        if first_word in ACTION_VERBS:
            action_verb_hits += 1

    if weak_hits:
        weaknesses.append(f"Found {weak_hits} bullet(s) using weak/passive phrasing (e.g. 'responsible for', 'worked on').")
        recommendations.append("Replace passive phrasing with strong action verbs (e.g. 'Led', 'Delivered', 'Optimized').")
        score -= min(15, weak_hits * 3)

    if all_bullets and quantified_hits / max(1, len(all_bullets)) < 0.3:
        weaknesses.append("Most bullet points lack quantifiable results (numbers, %, $, time saved, etc.).")
        recommendations.append("Quantify achievements wherever possible - e.g. 'Reduced processing time by 30%' instead of 'Improved processing time'.")
        score -= 10
    elif quantified_hits:
        strengths.append("Several bullet points include quantifiable, measurable results.")

    if all_bullets and action_verb_hits / max(1, len(all_bullets)) < 0.4:
        recommendations.append("Start more bullet points with strong action verbs rather than nouns or 'I'.")
        score -= 5
    elif action_verb_hits:
        strengths.append("Bullet points generally start with strong action verbs.")

    buzz_hits = [b for b in BUZZWORDS if b in cv.raw_text.lower()]
    if buzz_hits:
        weaknesses.append(f"Overused buzzwords detected: {', '.join(buzz_hits)}.")
        recommendations.append("Replace generic buzzwords with specific, evidence-backed accomplishments.")
        score -= 5

    word_count = len(cv.raw_text.split())
    if word_count > 900:
        weaknesses.append("CV is quite long - may exceed the ideal 1-2 page length.")
        recommendations.append("Trim to the most relevant, recent experience; aim for 1 page (early career) to 2 pages (senior).")
        score -= 5
    elif word_count < 150:
        weaknesses.append("CV content seems very sparse.")
        recommendations.append("Add more detail on responsibilities, achievements, and skills.")
        score -= 10

    score = max(0, min(100, score))
    if not recommendations:
        recommendations.append("Great job - only minor polish needed. Proofread once more for consistency in tense and formatting.")

    return Feedback(
        overall_score=score,
        strengths=strengths or ["CV contains the basic building blocks of a standard resume."],
        weaknesses=weaknesses or ["No major structural issues detected."],
        recommendations=recommendations,
    )


# Calls the Gemini API to generate feedback and rewritten content
def llm_feedback(cv: ParsedCV, model: str = "gemini-2.5-flash") -> Feedback:
    client = genai.Client(api_key=API_KEY)

    bullet_lines = [
        line.strip() for s in cv.sections for line in s.content
        if line.strip().startswith(("-", "•")) and not JOB_HEADER_RE.search(line)
    ]
    sections_overview = "\n".join(f"- {s.title}" for s in cv.sections)

    user_prompt = (
        f"NAME: {cv.contact.name}\n"
        f"CURRENT SUMMARY: {' '.join(cv.summary)}\n\n"
        f"SECTION TITLES PRESENT: \n{sections_overview}\n\n"
        f"BULLET LINES TO EVALUATE:\n" + "\n".join(bullet_lines) + "\n\n"
        f"Please analyze this CV and return the JSON feedback object as instructed."
    )

    response = client.models.generate_content(
        model=model,
        contents=user_prompt,
        config=types.GenerateContentConfig(
            system_instruction=LLM_SYSTEM_PROMPT,
            max_output_tokens=4000,
            response_mime_type="application/json",
        ),
    )

    text = response.text.strip()
    text = re.sub(r"^```json\s*|\s*```$", "", text)
    data = json.loads(text)

    return Feedback(
        overall_score=int(data.get("overall_score", 70)),
        strengths=data.get("strengths", []),
        weaknesses=data.get("weaknesses", []),
        recommendations=data.get("recommendations", []),
        rewritten_summary=data.get("rewritten_summary"),
        rewritten_bullets=data.get("rewritten_bullets"),
        raw_llm_response=text,
    )


# Tries the LLM path first, falls back to rule-based analysis on failure
def generate_feedback(cv: ParsedCV, force_rule_based: bool = False, model: str = "gemini-2.5-flash") -> Feedback:
    if not force_rule_based and API_KEY and "REPLACE" not in API_KEY:
        try:
            return llm_feedback(cv, model=model)
        except Exception as e:
            print(f"[warn] LLM feedback failed ({e}); falling back to rule-based analysis.", file=sys.stderr)
    return rule_based_feedback(cv)


# Local heuristic rewrite, used only when no LLM rewrite is available
def improve_bullet(line: str) -> str:
    text = line.strip("-• \t")
    low = text.lower()
    had_weak_phrase = False
    for phrase in WEAK_PHRASES:
        if phrase in low:
            idx = low.find(phrase)
            text = text[:idx] + text[idx + len(phrase):]
            text = text.strip(" ,.")
            had_weak_phrase = True
            break
    if text and text[0].islower():
        text = text[0].upper() + text[1:]
    if had_weak_phrase:
        first_word = re.sub(r"[^a-zA-Z]", "", text.split()[0]).lower() if text.split() else ""
        if first_word not in ACTION_VERBS and text:
            text = "Contributed to " + text[0].lower() + text[1:]
    return text


# Applies feedback to produce improved summary and bullet content
def build_optimized_content(cv: ParsedCV, feedback: Feedback) -> ParsedCV:
    optimized = ParsedCV(
        raw_text=cv.raw_text,
        contact=cv.contact,
        summary=list(cv.summary),
        sections=[CVSection(title=s.title, content=list(s.content)) for s in cv.sections],
    )

    if feedback.rewritten_summary:
        optimized.summary = [feedback.rewritten_summary]
    elif not cv.summary and cv.contact.name:
        optimized.summary = ["Motivated professional with a track record of delivering results."]

    rewrite_lookup = {k.strip(): v for k, v in (feedback.rewritten_bullets or {}).items()}

    for section in optimized.sections:
        new_content = []
        for line in section.content:
            stripped = line.strip()
            is_bullet = stripped.startswith(("-", "•")) and not JOB_HEADER_RE.search(stripped)
            if not is_bullet:
                new_content.append(line)
            elif stripped in rewrite_lookup:
                new_content.append(rewrite_lookup[stripped])
            else:
                new_content.append(improve_bullet(line))
        section.content = new_content

    return optimized

    return optimized


# Adds a styled section heading with an underline rule
def _add_section_heading(document, text: str):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = ACCENT_COLOR
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '1', qn('w:color'): 'A6A6A6'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# Renders the optimized CV into a formatted .docx file
def render_docx(cv: ParsedCV, output_path: str, feedback: Optional[Feedback] = None):
    document = docx.Document()

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    section = document.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    name_p = document.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv.contact.name or "Your Name")
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = ACCENT_COLOR
    name_p.space_after = Pt(2)

    contact_bits = [b for b in [cv.contact.location, cv.contact.email, cv.contact.phone, cv.contact.linkedin] if b]
    if contact_bits:
        contact_p = document.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact_p.add_run("  |  ".join(contact_bits))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        contact_p.space_after = Pt(10)

    divider = document.add_paragraph()
    divider.paragraph_format.space_after = Pt(6)
    pPr = divider._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '8', qn('w:space'): '1', qn('w:color'): '1F4E79'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)

    if cv.summary:
        _add_section_heading(document, "PROFESSIONAL SUMMARY")
        p = document.add_paragraph(" ".join(cv.summary))
        p.paragraph_format.space_after = Pt(8)

    for sec in cv.sections:
        _add_section_heading(document, sec.title.upper())
        for line in sec.content:
            clean = line.strip()
            if not clean:
                continue
            if JOB_HEADER_RE.search(clean):
                parts = re.split(r"\s*[|·]\s*", clean)
                p = document.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                run = p.add_run(parts[0])
                run.bold = True
                run.font.size = Pt(11)
                if len(parts) > 1:
                    tab_run = p.add_run("\t" + "  |  ".join(parts[1:]))
                    tab_run.italic = True
                    tab_run.font.size = Pt(10)
                    tab_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            elif clean.startswith(("-", "•")):
                bullet_text = clean.lstrip("-•").strip()
                p = document.add_paragraph(bullet_text, style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
            else:
                p = document.add_paragraph(clean)
                p.paragraph_format.space_after = Pt(6)

    if feedback is not None:
        document.add_page_break()
        _add_section_heading(document, "CV REVIEW SUMMARY (FOR REFERENCE - REMOVE BEFORE SENDING)")
        p = document.add_paragraph(f"Overall Score: {feedback.overall_score}/100")
        p.runs[0].bold = True

        document.add_paragraph("Strengths:", style="Heading 3")
        for s in feedback.strengths:
            document.add_paragraph(s, style="List Bullet")

        document.add_paragraph("Areas Improved / To Watch:", style="Heading 3")
        for w in feedback.weaknesses:
            document.add_paragraph(w, style="List Bullet")

        document.add_paragraph("Recommendations Applied:", style="Heading 3")
        for r in feedback.recommendations:
            document.add_paragraph(r, style="List Bullet")

    document.save(output_path)


# Converts the .docx to PDF using LibreOffice, if available
def convert_docx_to_pdf(docx_path: str) -> Optional[str]:
    out_dir = str(Path(docx_path).parent)
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            check=True, capture_output=True, timeout=120,
        )
        pdf_path = str(Path(docx_path).with_suffix(".pdf"))
        return pdf_path if Path(pdf_path).exists() else None
    except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired) as e:
        print(f"[warn] PDF conversion skipped (LibreOffice not available: {e}).", file=sys.stderr)
        return None


# Builds the Markdown feedback report
def render_feedback_markdown(feedback: Feedback, cv: ParsedCV) -> str:
    lines = [
        "# CV Feedback Report",
        f"**Candidate:** {cv.contact.name or 'N/A'}",
        f"**Overall Score:** {feedback.overall_score}/100",
        "",
        "## Strengths",
    ]
    lines += [f"- {s}" for s in feedback.strengths] or ["- None identified"]
    lines += ["", "## Weaknesses"]
    lines += [f"- {w}" for w in feedback.weaknesses] or ["- None identified"]
    lines += ["", "## Recommendations"]
    lines += [f"- {r}" for r in feedback.recommendations] or ["- None"]

    if feedback.rewritten_summary:
        lines += ["", "## Suggested Summary Rewrite", "", feedback.rewritten_summary]

    if feedback.rewritten_bullets:
        lines += ["", "## Suggested Bullet Rewrites"]
        for section, bullets in feedback.rewritten_bullets.items():
            lines += [f"\n**{section}**"]
            lines += [f"- {b}" for b in bullets]

    return "\n".join(lines)


# Command-line entry point
def main():
    parser = argparse.ArgumentParser(description="LLM-Powered CV Feedback and Optimization tool.")
    parser.add_argument("--input", "-i", required=True, help="Path to the input CV (.txt, .docx, .pdf)")
    parser.add_argument("--output", "-o", default="optimized_cv.docx", help="Path for the optimized CV (.docx)")
    parser.add_argument("--feedback-out", "-f", default="feedback_report.md", help="Path for the feedback report (Markdown)")
    parser.add_argument("--pdf", action="store_true", help="Also export the optimized CV as PDF (requires LibreOffice)")
    parser.add_argument("--rule-based", action="store_true", help="Force the offline rule-based analyzer")
    parser.add_argument("--model", default="gemini-2.5-flash", help="Gemini model to use for LLM feedback")
    parser.add_argument("--include-review-page", action="store_true", help="Append feedback summary as an extra page in the docx")
    args = parser.parse_args()

    print(f"[1/5] Reading CV from: {args.input}")
    raw_text = read_cv_file(args.input)

    print("[2/5] Parsing CV structure...")
    cv = parse_cv(raw_text)

    print("[3/5] Generating feedback...")
    feedback = generate_feedback(cv, force_rule_based=args.rule_based, model=args.model)
    print(f"      -> Overall Score: {feedback.overall_score}/100")

    print(f"[4/5] Writing feedback report to: {args.feedback_out}")
    Path(args.feedback_out).write_text(render_feedback_markdown(feedback, cv), encoding="utf-8")

    print("[5/5] Building optimized CV...")
    optimized_cv = build_optimized_content(cv, feedback)
    render_docx(optimized_cv, args.output, feedback=feedback if args.include_review_page else None)
    print(f"      -> Optimized CV saved to: {args.output}")

    if args.pdf:
        pdf_path = convert_docx_to_pdf(args.output)
        if pdf_path:
            print(f"      -> PDF version saved to: {pdf_path}")

    print("\nDone!")


if __name__ == "__main__":
    main()